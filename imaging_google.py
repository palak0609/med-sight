import os
from PIL import Image as PILImage
from agno.agent import Agent
from agno.models.google import Gemini
import streamlit as st
# from agno.tools.duckduckgo import DuckDuckGoTools  # Commented out - internet search disabled
from agno.media import Image as AgnoImage
from docx import Document
from docx.shared import Inches
import io
import re
from dotenv import load_dotenv
import pydicom
import numpy as np

# Load environment variables from a .env file (for local/dev) and OS env (for prod)
load_dotenv()
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")

with st.sidebar:
    st.title("ℹ️ Configuration")
    st.info(
        "This tool provides AI-powered analysis of medical imaging data using "
        "advanced computer vision and radiological expertise."
    )
    st.warning(
        "⚠DISCLAIMER: This tool is for educational and informational purposes only. "
        "All analyses should be reviewed by qualified healthcare professionals. "
        "Do not make medical decisions based solely on this analysis."
    )
    if GOOGLE_API_KEY:
        st.success("Workspace ready!")
    else:
        st.error(
            "Google API key is not configured on the server. "
            "Please set GOOGLE_API_KEY in your environment or .env file."
        )

medical_agent = Agent(
    model=Gemini(
        id="gemini-2.0-flash",
        api_key=GOOGLE_API_KEY
    ),
    # tools=[DuckDuckGoTools()],  # Commented out - internet search disabled
    tools=[],  # No tools - internet search disabled
    markdown=True
) if GOOGLE_API_KEY else None

if not medical_agent:
    st.warning("Service is not available because the API key is not configured.")

# Medical Analysis Query
query = """
You are a highly skilled medical expert specializing in diagnostic imaging and analytical interpretation of medical graphs. Analyze the provided medical image or graph comprehensively, solely based on the visual data available, and structure your analysis clearly under the following headings:

### 1. Image/Graph Type & Anatomical Region
- Specify the imaging or graph modality clearly (X-ray, MRI, CT, Ultrasound, ECG, EEG, etc.)
- Identify the anatomical region or physiological system depicted
- Comment on the technical quality, clarity, and adequacy of the data provided


### 2. Detailed Observations & Key Findings
- Systematically list primary visual or numerical observations
- Clearly describe any abnormalities, anomalies, or deviations from normal reference ranges
- Include precise measurements, numerical values, or densities as applicable
- Clearly describe the location, size, shape, characteristics, or pattern of abnormalities
- Rate severity clearly as Normal, Mild, Moderate, or Severe

### 3. Diagnostic Assessment
- State your primary diagnosis clearly and confidently based solely on visual or graphical evidence
- List possible differential diagnoses ranked by likelihood, supported by the provided data
- Highlight specific visual or numerical evidence underpinning each diagnosis
- Clearly flag any critical, urgent, or emergent findings requiring immediate medical attention

### 4. Patient-Friendly Explanation
- Clearly and simply explain your findings without medical jargon
- Provide definitions or analogies to help the patient understand the significance
- Address potential patient concerns about severity, prognosis, or immediate next steps
- Suggest general recommendations or follow-up steps clearly understandable by a non-medical individual

Ensure your analysis remains precise, thorough, and clear, making it fully applicable to multimodal medical data including both images and graphical presentations. Consider that your interpretation may need to stand independently without additional patient history or symptomatic context.
"""

st.title("🏥 Medical Imaging Diagnosis Agent")
st.write("Upload a medical image for professional analysis")

# Create containers for better organization
upload_container = st.container()
image_container = st.container()
analysis_container = st.container()

with upload_container:
    uploaded_file = st.file_uploader(
        "Upload Medical Image",
        type=["jpg", "jpeg", "png", "dicom"],
        help="Supported formats: JPG, JPEG, PNG, DICOM"
    )

if uploaded_file is not None:
    with image_container:
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            # Handle standard images and DICOM files
            file_name = uploaded_file.name.lower()
            if file_name.endswith(".dcm") or file_name.endswith(".dicom"):
                try:
                    ds = pydicom.dcmread(uploaded_file)
                    pixel_array = ds.pixel_array.astype(float)
                    # Normalize to 0-255 for display
                    pixel_array -= pixel_array.min()
                    if pixel_array.max() > 0:
                        pixel_array /= pixel_array.max()
                    pixel_array = (pixel_array * 255).astype("uint8")
                    image = PILImage.fromarray(pixel_array)
                except Exception as e:
                    st.error(f"Failed to read DICOM file: {e}")
                    image = None
            else:
                image = PILImage.open(uploaded_file)

            if image is None:
                st.stop()
            width, height = image.size
            aspect_ratio = width / height
            new_width = 500
            new_height = int(new_width / aspect_ratio)
            resized_image = image.resize((new_width, new_height))
            
            st.image(
                resized_image,
                caption="Uploaded Medical Image",
                use_container_width=True
            )
            
            analyze_button = st.button(
                "🔍 Analyze Image",
                type="primary",
                use_container_width=True
            )
    
    with analysis_container:
        if analyze_button:
            with st.spinner("🔄 Analyzing image... Please wait."):
                try:
                    temp_path = "temp_resized_image.png"
                    resized_image.save(temp_path)
                    
                    # Create AgnoImage object
                    agno_image = AgnoImage(filepath=temp_path)  # Adjust if constructor differs
                    
                    # Run analysis
                    response = medical_agent.run(query, images=[agno_image])
                    st.markdown("### 📋 Analysis Results")
                    st.markdown("---")
                    st.markdown(response.content)
                    st.markdown("---")
                    st.caption(
                        "Note: This analysis is generated by AI and should be reviewed by "
                        "a qualified healthcare professional."
                    )

                    # --- DOCX Export Functionality ---
                    def create_docx(analysis_markdown, image_pil):
                        doc = Document()
                        doc.add_heading('Medical Imaging Analysis', 0)
                        # Add image at the top
                        doc.add_heading('Uploaded Image', level=1)
                        img_stream = io.BytesIO()
                        image_pil.save(img_stream, format='PNG')
                        img_stream.seek(0)
                        doc.add_picture(img_stream, width=Inches(5))
                        # Add analysis after image
                        bold_pattern = re.compile(r'(\*\*|__)(.+?)\1')
                        for line in analysis_markdown.split('\n'):
                            if line.startswith('#'):
                                doc.add_heading(line.replace('#', '').strip(), level=1)
                            elif line.strip() == '---':
                                doc.add_page_break()
                            elif line.strip():
                                # Handle bold markdown
                                p = doc.add_paragraph()
                                last_idx = 0
                                for m in bold_pattern.finditer(line):
                                    # Add text before bold
                                    if m.start() > last_idx:
                                        p.add_run(line[last_idx:m.start()])
                                    # Add bold text
                                    p.add_run(m.group(2)).bold = True
                                    last_idx = m.end()
                                # Add any remaining text
                                if last_idx < len(line):
                                    p.add_run(line[last_idx:])
                        return doc

                    # Store analysis and image for download
                    if 'analysis_result' not in st.session_state:
                        st.session_state['analysis_result'] = None
                    if 'analysis_image' not in st.session_state:
                        st.session_state['analysis_image'] = None
                    st.session_state['analysis_result'] = response.content
                    st.session_state['analysis_image'] = resized_image

                    # Download button
                    def get_docx_download():
                        doc = create_docx(st.session_state['analysis_result'], st.session_state['analysis_image'])
                        docx_io = io.BytesIO()
                        doc.save(docx_io)
                        docx_io.seek(0)
                        return docx_io

                    st.download_button(
                        label="⬇️ Download Analysis as DOCX",
                        data=get_docx_download(),
                        file_name="medical_image_analysis.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    # --- End DOCX Export ---
                except Exception as e:
                    st.error(f"Analysis error: {e}")
else:
    st.info("👆 Please upload a medical image to begin analysis")
